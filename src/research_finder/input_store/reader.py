"""Read plain text from the file formats we accept in `input/raw/`."""
from __future__ import annotations

from pathlib import Path

SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx", ".pdf"}


def read_text(path: Path) -> str:
    """Extract plain text from a file. Raises ValueError on unsupported extensions."""
    ext = path.suffix.lower()
    if ext in {".md", ".txt"}:
        return path.read_text(encoding="utf-8")
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    raise ValueError(f"Unsupported file extension: {ext} (supported: {sorted(SUPPORTED_EXTENSIONS)})")


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    import fitz

    with fitz.open(str(path)) as doc:
        return "\n".join(page.get_text() for page in doc)
